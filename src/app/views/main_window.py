from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from datetime import date
from .const import COMPANY, ADDRESS, CITY, TAX_CODE, VAT_ID, EMAIL
import uuid
import sys
import subprocess
import os
from pathlib import Path
from app.views.custom_form_field import FormField
from app.config_manager import ConfigManager

def generate_estimate():
    document = Document()
    header = document.sections[0].header
    document.sections[0].header_distance = Inches(1.5)
    table = header.add_table(rows=1, cols=3, width=Inches(6.5))  # Full page width
    table.autofit = False

    # col1
    left_col = table.cell(0, 0)
    left_col.width = Inches(2.5)
    left_para = left_col.paragraphs[0]
    left_para.add_run("Informazioni cliente")

    left_col.paragraphs[0].paragraph_format.border_bottom = True
    left_col.paragraphs[0].paragraph_format.border_top = True
    left_col.paragraphs[0].paragraph_format.border_left = True
    left_col.paragraphs[0].paragraph_format.border_right = True

    # col2
    middle_col = table.cell(0, 1)
    middle_col.width = Inches(1.5)
    middle_para = middle_col.paragraphs[0]
    middle_para.add_run("\n\n")
    middle_para.add_run("Preventivo").bold = True

    # col3
    right_col = table.cell(0, 2)
    right_col.width = Inches(2.5)
    right_para = right_col.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    right_para.add_run(COMPANY + "\n")
    right_para.add_run(ADDRESS + "\n")
    right_para.add_run(CITY + "\n")
    right_para.add_run(TAX_CODE + "\n")
    right_para.add_run(VAT_ID + "\n")
    right_para.add_run(EMAIL + "\n")

    date_sig = date.today().strftime('%d-%m-%Y')  # file signature
    uuid_sig = uuid.uuid4().hex[:8]
    document.add_paragraph("Corpo documento...")

    downloads_dir = Path.home() / "Downloads"
    filename = "Preventivo_{date}_{uuid}.docx".format(date=date_sig, uuid=uuid_sig)

    directory = downloads_dir / filename
    document.save(str(directory))
    open_file(directory)

def open_file(filepath):
    try:
        if sys.platform == "win32":
            os.startfile(filepath)
        elif sys.platform == "darwin":
            subprocess.run(["open", filepath])
        else:  # Linux
            subprocess.run(["xdg-open", filepath])
    except Exception as e:
        print(f"Error opening file: {e}")

class MainWindow(QMainWindow):
    def show_settings_page(self):
        # Redirects to settings page
        self.stacked_widget.setCurrentIndex(1)
        self.setFixedSize(QSize(600, 700))

    def show_main_page(self):
        # Redirect to main page
        self.stacked_widget.setCurrentIndex(0)
        self.setFixedSize(QSize(500, 400))

    def create_settings_page(self):

        widget = QWidget()
        layout = QVBoxLayout(widget)
        self.setFixedSize(QSize(500, 800))

        self.company_field = FormField("Nome Azienda")
        self.address_field = FormField("Indirizzo", "Via Roma, 28, TO")
        self.city_field = FormField("Città", "Roma")
        self.tax_code_field = FormField("Codice Fiscale", "NTAPQL61S01D568L", "[A-Z0-9]{0,16}")
        self.vat_field = FormField("Partita IVA", "17556120099")
        self.email_field = FormField("Email", "example@gmail.com")

        button_layout = QHBoxLayout()

        save_btn = QPushButton("Salva")
        save_btn.clicked.connect(self.save_settings)

        back_btn = QPushButton("Indietro")
        back_btn.clicked.connect(self.show_main_page)

        layout.addWidget(self.company_field)
        layout.addWidget(self.address_field)
        layout.addWidget(self.city_field)
        layout.addWidget(self.tax_code_field)
        layout.addWidget(self.vat_field)
        layout.addWidget(self.email_field)

        button_layout.addWidget(save_btn)
        button_layout.addWidget(back_btn)
        layout.addLayout(button_layout)

        return widget

    def save_settings(self):

        params = {
            "company_name": self.company_field.input.text(),
            "address": self.address_field.input.text(),
            "city": self.city_field.input.text(),
            "tax_code": self.tax_code_field.input.text(),
            "vat": self.vat_field.input.text(),
            "email": self.email_field.input.text()
        }

        if self.config_manager.save_config(params):
            QMessageBox.information(self, "Successo", "Dati salvati con successo!")
            self.current_config = params
            self.show_main_page()
        else:
            QMessageBox.critical(self, "Errore", "Impossibile salvare i dati!")

        self.show_main_page()

    def __init__(self, *args, **kwargs):
        super(MainWindow, self).__init__(*args, **kwargs)
        self.config_manager = ConfigManager()

        self.stacked_widget = QStackedWidget()
        self.setCentralWidget(self.stacked_widget)

        central_widget = QWidget()
        layout = QVBoxLayout(central_widget)

        self.setWindowTitle("Generatore template preventivo")

        settings_btn = QPushButton("Modifica dati personali")
        settings_btn.clicked.connect(self.show_settings_page)
        settings_btn.setFixedHeight(100)

        btn = QPushButton("Genera Preventivo")
        btn.clicked.connect(generate_estimate)
        btn.setFixedHeight(100)

        layout.addWidget(settings_btn)
        layout.addWidget(btn)

        self.stacked_widget.addWidget(central_widget)

        settings_page = self.create_settings_page()
        self.stacked_widget.addWidget(settings_page)

        self.setFixedSize(QSize(500, 400))
